#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Simple RAG Training Script

This single script handles everything:
- Auto-detects file types (DOCX, PDF, TXT)
- Converts files to text
- Trains the RAG system
- Shows progress and statistics

Usage:
    python train_rag.py "your-file.docx"
    python train_rag.py "your-file.pdf" 
    python train_rag.py "your-file.txt"
    python train_rag.py rag-training-data  # Train all files in directory
    python train_rag.py  # Interactive mode
"""

import asyncio
import argparse
import os
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional

# Fix Windows console encoding
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# Add the app directory to the path
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from dotenv import load_dotenv
load_dotenv()

# Try to import required libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from app.ai.embedding_service import get_embedding_service
    from app.ai.vector_storage import vector_storage
    from app.database.supabase import get_supabase_client
    RAG_AVAILABLE = True
except ImportError as e:
    print(f"âŒ RAG components not available: {e}")
    RAG_AVAILABLE = False


class SimpleRAGTrainer:
    """Simple RAG trainer that handles everything"""
    
    def __init__(self):
        if not RAG_AVAILABLE:
            raise Exception("RAG components not available. Check your setup.")
        
        self.embedding_service = get_embedding_service()
        self.vector_storage = vector_storage
        self.supabase = get_supabase_client()
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension"""
        ext = Path(file_path).suffix.lower()
        if ext == '.docx':
            return 'docx'
        elif ext == '.pdf':
            return 'pdf'
        elif ext in ['.txt', '.md']:
            return 'text'
        else:
            return 'unknown'
    
    def convert_to_text(self, file_path: str) -> str:
        """Convert any supported file to text"""
        file_type = self.detect_file_type(file_path)
        
        print(f"ğŸ“„ Converting {file_type.upper()} file: {Path(file_path).name}")
        
        if file_type == 'docx':
            return self._convert_docx(file_path)
        elif file_type == 'pdf':
            return self._convert_pdf(file_path)
        elif file_type == 'text':
            return self._read_text(file_path)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
    
    def _convert_docx(self, file_path: str) -> str:
        """Convert DOCX to text"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error converting DOCX: {e}")
    
    def _convert_pdf(self, file_path: str) -> str:
        """Convert PDF to text"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 not installed. Run: pip install PyPDF2")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text.strip())
                
                return "\n\n".join(text_content)
                
        except Exception as e:
            raise Exception(f"Error converting PDF: {e}")
    
    def _read_text(self, file_path: str) -> str:
        """Read text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error reading text file: {e}")
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                search_start = max(start + chunk_size - 100, start)
                for i in range(end - 1, search_start, -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    async def add_knowledge_chunk(self, content: str, source: str, knowledge_type: str = "global") -> bool:
        """Add a knowledge chunk to RAG"""
        try:
            # Generate embedding
            embedding = await self.embedding_service.generate_embedding(content)
            
            # Extract category from source filename
            source_name = Path(source).stem
            category = "general"
            if "story" in source_name.lower():
                category = "storytelling"
            elif "character" in source_name.lower():
                category = "character"
            elif "plot" in source_name.lower():
                category = "plot"
            elif "dialogue" in source_name.lower():
                category = "dialogue"
            
            # Store in vector database using existing function
            result = await self.vector_storage.store_global_knowledge(
                category=category,
                pattern_type="knowledge_chunk",
                embedding=embedding,
                example_text=content,
                description=f"Knowledge chunk from {source_name}",
                quality_score=0.8,
                tags=[source_name, "training"]
            )
            
            return result is not None
            
        except Exception as e:
            print(f"âŒ Error adding chunk: {e}")
            return False
    
    async def train_from_file(self, file_path: str, knowledge_type: str = "global") -> bool:
        """Train RAG from any supported file"""
        try:
            # Convert file to text
            text_content = self.convert_to_text(file_path)
            
            print(f"ğŸ“Š Content length: {len(text_content)} characters")
            
            # Split into chunks
            chunks = self.chunk_text(text_content)
            print(f"ğŸ“¦ Split into {len(chunks)} chunks")
            
            # Add each chunk to RAG
            success_count = 0
            for i, chunk in enumerate(chunks):
                print(f"ğŸ”„ Processing chunk {i+1}/{len(chunks)}")
                
                metadata = {
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_name": Path(file_path).name
                }
                
                success = await self.add_knowledge_chunk(
                    content=chunk,
                    source=file_path,
                    knowledge_type=knowledge_type
                )
                
                if success:
                    success_count += 1
                
                # Small delay to avoid rate limiting
                await asyncio.sleep(0.1)
            
            print(f"ğŸ“ˆ Successfully processed {success_count}/{len(chunks)} chunks")
            return success_count > 0
            
        except Exception as e:
            print(f"âŒ Error training from file: {e}")
            return False
    
    async def show_statistics(self):
        """Show knowledge base statistics"""
        try:
            result = self.supabase.table('global_knowledge').select('*').execute()
            
            if result.data:
                total_chunks = len(result.data)
                
                print(f"\nğŸ“Š Knowledge Base Statistics:")
                print(f"   Total knowledge chunks: {total_chunks}")
                
                # Show categories
                categories = {}
                for item in result.data:
                    category = item.get('category', 'Unknown')
                    categories[category] = categories.get(category, 0) + 1
                
                print(f"\nğŸ“ Categories ({len(categories)}):")
                for category, count in sorted(categories.items()):
                    print(f"   - {category}: {count} chunks")
                
                # Show pattern types
                pattern_types = {}
                for item in result.data:
                    pattern_type = item.get('pattern_type', 'Unknown')
                    pattern_types[pattern_type] = pattern_types.get(pattern_type, 0) + 1
                
                print(f"\nğŸ” Pattern Types ({len(pattern_types)}):")
                for pattern_type, count in sorted(pattern_types.items()):
                    print(f"   - {pattern_type}: {count} chunks")
            else:
                print("ğŸ“Š No knowledge chunks found in database")
                
        except Exception as e:
            print(f"âŒ Error retrieving statistics: {e}")
    
    async def interactive_mode(self):
        """Interactive mode for adding knowledge"""
        print("ğŸ“ RAG Training - Interactive Mode")
        print("=" * 40)
        
        while True:
            print("\nOptions:")
            print("1. Train from file")
            print("2. Add text content")
            print("3. View statistics")
            print("4. Exit")
            
            choice = input("\nEnter your choice (1-4): ").strip()
            
            if choice == '1':
                file_path = input("Enter file path: ").strip()
                if os.path.exists(file_path):
                    knowledge_type = input("Knowledge type (global/user) [global]: ").strip() or "global"
                    success = await self.train_from_file(file_path, knowledge_type)
                    if success:
                        print("âœ… File trained successfully!")
                    else:
                        print("âŒ Failed to train file")
                else:
                    print("âŒ File not found")
            
            elif choice == '2':
                source = input("Source/Title: ").strip() or "Interactive Input"
                print("Enter your content (press Ctrl+D or Ctrl+Z when done):")
                try:
                    content = ""
                    while True:
                        line = input()
                        content += line + "\n"
                except EOFError:
                    pass
                
                if content.strip():
                    knowledge_type = input("Knowledge type (global/user) [global]: ").strip() or "global"
                    success = await self.add_knowledge_chunk(content.strip(), source, knowledge_type)
                    if success:
                        print("âœ… Content added successfully!")
                    else:
                        print("âŒ Failed to add content")
                else:
                    print("âŒ No content provided")
            
            elif choice == '3':
                await self.show_statistics()
            
            elif choice == '4':
                print("ğŸ‘‹ Goodbye!")
                break
            
            else:
                print("âŒ Invalid choice")
    
    async def train_directory(self, directory_path: str, knowledge_type: str = "global") -> Dict[str, bool]:
        """Train RAG from all supported files in a directory"""
        directory = Path(directory_path)
        if not directory.exists() or not directory.is_dir():
            raise Exception(f"Directory not found: {directory_path}")
        
        # Find all supported files
        supported_extensions = ['.pdf', '.docx', '.txt', '.md']
        files = []
        for ext in supported_extensions:
            files.extend(directory.glob(f'*{ext}'))
        
        if not files:
            print(f"âš ï¸ No supported files found in {directory_path}")
            return {}
        
        print(f"ğŸ“ Found {len(files)} files in {directory_path}")
        print("=" * 60)
        
        results = {}
        for i, file_path in enumerate(files, 1):
            print(f"\n[{i}/{len(files)}] Processing: {file_path.name}")
            print("-" * 60)
            success = await self.train_from_file(str(file_path), knowledge_type)
            results[str(file_path)] = success
            
            if success:
                print(f"âœ… {file_path.name} - Training completed!")
            else:
                print(f"âŒ {file_path.name} - Training failed!")
        
        return results

async def main():
    """Main function"""
    parser = argparse.ArgumentParser(description='Simple RAG Training Script')
    parser.add_argument('file', nargs='?', help='File or directory to train with (DOCX, PDF, TXT)')
    parser.add_argument('--type', default='global', help='Knowledge type (global/user)')
    parser.add_argument('--interactive', action='store_true', help='Run in interactive mode')
    
    args = parser.parse_args()
    
    # Check environment
    required_vars = ['OPENAI_API_KEY', 'SUPABASE_URL', 'SUPABASE_SERVICE_ROLE_KEY']
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    
    if missing_vars:
        print(f"âŒ Missing environment variables: {', '.join(missing_vars)}")
        print("Please check your .env file")
        return
    
    if not RAG_AVAILABLE:
        print("âŒ RAG components not available")
        print("Make sure you're in the backend directory and all dependencies are installed")
        return
    
    try:
        trainer = SimpleRAGTrainer()
        
        if args.interactive or not args.file:
            # Interactive mode
            await trainer.interactive_mode()
        else:
            # Check if it's a file or directory
            path = Path(args.file)
            if not path.exists():
                print(f"âŒ Path not found: {args.file}")
                return
            
            if path.is_dir():
                # Train all files in directory
                print(f"ğŸš€ Training RAG from directory: {path.name}")
                print("=" * 60)
                results = await trainer.train_directory(str(path), args.type)
                
                # Summary
                print("\n" + "=" * 60)
                print("ğŸ“Š Training Summary:")
                print("=" * 60)
                successful = sum(1 for success in results.values() if success)
                failed = len(results) - successful
                print(f"âœ… Successful: {successful}/{len(results)}")
                print(f"âŒ Failed: {failed}/{len(results)}")
                
                if successful > 0:
                    await trainer.show_statistics()
            else:
                # Train from single file
                print(f"ğŸš€ Training RAG with: {path.name}")
                success = await trainer.train_from_file(str(path), args.type)
                
                if success:
                    print("âœ… RAG training completed successfully!")
                    await trainer.show_statistics()
                else:
                    print("âŒ RAG training failed")
    
    except Exception as e:
        print(f"âŒ Error: {e}")
        import traceback
        traceback.print_exc()
        print("\nğŸ’¡ Make sure you have installed required dependencies:")
        print("   pip install python-docx PyPDF2")


if __name__ == "__main__":
    asyncio.run(main())
